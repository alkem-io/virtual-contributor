"""Extract plain text from link contribution bodies.

Supported formats (explicit allowlist): PDF, DOCX, XLSX, plain text, HTML.
Anything else is skipped — there's no point feeding binary or opaque
content to the chunker/embedder.  The fetch layer enforces a hard byte
cap; chunking + content-hash change detection handle incremental
re-ingest, so we deliberately do *not* truncate the extracted text.
"""

from __future__ import annotations

import io
import logging
import re

logger = logging.getLogger(__name__)

_WHITESPACE_RE = re.compile(r"[ \t]+")
_EXCESS_NEWLINES_RE = re.compile(r"\n[ \t]*\n[ \t]*\n+")

# MIME tokens we're willing to decode.  Check with "in" against a
# normalised content_type string — covers e.g. "application/pdf",
# "application/pdf; charset=binary", etc.
_MIME_KIND = {
    "pdf": "pdf",
    "vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "html": "html",
    "xml": "html",
    "text/plain": "text",
    "text/csv": "text",
    "text/markdown": "text",
    "json": "text",
}

# Fallback: sniff the first few bytes when the server didn't set a
# useful content-type.  ZIP-based Office formats all start with "PK".
_MAGIC = [
    (b"%pdf", "pdf"),
    (b"<!doc", "html"),
    (b"<html", "html"),
    (b"<?xml", "html"),
]


def extract_text(body: bytes, content_type: str) -> str | None:
    """Return extracted plain text, or ``None`` if the format isn't supported."""
    if not body:
        return None

    kind = _detect_kind(body, content_type)
    if kind is None:
        return None

    try:
        if kind == "pdf":
            text = _extract_pdf(body)
        elif kind == "docx":
            text = _extract_docx(body)
        elif kind == "xlsx":
            text = _extract_xlsx(body)
        elif kind == "docx_or_xlsx":
            # ZIP container — try DOCX first, fall back to XLSX.
            try:
                text = _extract_docx(body)
            except Exception:
                text = _extract_xlsx(body)
        elif kind == "html":
            text = _extract_html(body)
        else:
            text = body.decode("utf-8", errors="replace")
    except Exception as exc:
        logger.warning("Link text extraction failed (%s): %s", kind, exc)
        return None

    text = _normalise(text)
    return text or None


def _detect_kind(body: bytes, content_type: str) -> str | None:
    ct = (content_type or "").lower()
    for token, kind in _MIME_KIND.items():
        if token in ct:
            return kind
    head = body[:32]
    head_lower = head.lower()
    for signature, kind in _MAGIC:
        if head_lower.startswith(signature):
            return kind
    # ZIP magic (PK\x03\x04) may be DOCX or XLSX — we can't tell without
    # peeking inside.  Try DOCX first, then XLSX; both extractors are
    # safe (they raise on mismatch).
    if head.startswith(b"PK\x03\x04"):
        return "docx_or_xlsx"
    return None


def _extract_pdf(body: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(body))
    pages: list[str] = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        if page_text:
            pages.append(page_text)
    return "\n\n".join(pages)


def _extract_docx(body: bytes) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(body))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(body: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(body), data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"## {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [
                "" if v is None else str(v)
                for v in row
            ]
            if any(c.strip() for c in cells):
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


def _extract_html(body: bytes) -> str:
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        return body.decode("utf-8", errors="replace")

    soup = BeautifulSoup(body, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text(separator="\n")


def _normalise(text: str) -> str:
    text = _WHITESPACE_RE.sub(" ", text)
    text = _EXCESS_NEWLINES_RE.sub("\n\n", text)
    return text.strip()
